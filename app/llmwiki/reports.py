"""
reports.py — Grounded report generation pipeline.

Flow:
1. Load a ReportTemplate (column definitions, scope slugs).
2. For each page in scope, call LLM to extract the column values.
3. Assemble ExtractedRow objects.
4. Export as XLSX, DOCX, or PDF.
"""
from __future__ import annotations

import io
import json
import uuid
from datetime import UTC, datetime
from html import escape
from typing import Any

from sqlalchemy.orm import Session

from app.db.models import ReportJob, ReportTemplate
from app.llmwiki.storage import WikiStore
from app.llmwiki.text import safe_format, trim_to_chars
from app.schemas.workspace import ExtractedCell, ExtractedRow, ReportColumnDef
from app.services.llm_factory import JsonLlm


# ---------------------------------------------------------------------------
# LLM Extraction Prompt
# ---------------------------------------------------------------------------

EXTRACTION_PROMPT = """\
You are a structured data extractor. Given a wiki page and a list of extraction columns,
return a JSON object with one key per column (using the column key as the JSON key).

Each value should be an object with:
  - value: the extracted value as a string (or "" if not found).
    CRITICAL FORMATTING REQUIREMENT: If the extracted value contains lists, multiple items, qualifications, certifications, dates, or structured information, DO NOT write it in a single dense paragraph or run-on sentence. You MUST format it beautifully and cleanly using markdown lists (e.g. using bullet points with "-", numbered lists, bold text keywords, or simple markdown tables where appropriate) so that it renders in a highly structured, professional way in documents.
  - confidence: float 0.0–1.0 (1.0 = explicit statement, 0.5 = inferred, 0.0 = not found)
  - quote: the exact sentence from the page that supports this value (or null)

Return ONLY the JSON. Do NOT invent information not present in the page.

Columns to extract:
{columns_spec}

Wiki page title: {title}
Wiki page content:
{content}
"""

KNOWLEDGE_EXTRACTION_PROMPT = """\
You are an intelligent knowledge base assistant. Answer the following questions/extraction columns using your own internal knowledge. Do NOT look up or expect any wiki documents.
Return a JSON object with one key per column (using the column key as the JSON key).

Each value should be an object with:
  - value: your answer as a string (or "" if you do not know).
    CRITICAL FORMATTING REQUIREMENT: If your answer contains lists, multiple items, qualifications, certifications, dates, or structured information, DO NOT write it in a single dense paragraph or run-on sentence. You MUST format it beautifully and cleanly using markdown lists (e.g. using bullet points with "-", numbered lists, bold text keywords, or simple markdown tables where appropriate) so that it renders in a highly structured, professional way in documents.
  - confidence: float 0.0–1.0 (1.0 = highly certain, 0.5 = uncertain, 0.0 = completely unknown)
  - quote: null (since you are answering from your own knowledge)

Columns to answer:
{columns_spec}
"""


# ---------------------------------------------------------------------------
# Extractor
# ---------------------------------------------------------------------------


class ReportExtractor:
    def __init__(self, store: WikiStore, llm: JsonLlm | None = None):
        self.store = store
        self.llm = llm

    @staticmethod
    def _normalise_cell(raw: Any, *, slug: str) -> ExtractedCell:
        if isinstance(raw, dict):
            confidence = raw.get("confidence", 0.0)
            try:
                confidence_float = float(confidence)
            except (TypeError, ValueError):
                confidence_float = 0.0
            confidence_float = max(0.0, min(1.0, confidence_float))
            return ExtractedCell(
                value=str(raw.get("value", "") or ""),
                confidence=confidence_float,
                source_slug=slug,
                quote=str(raw.get("quote", "") or "") or None,
            )
        return ExtractedCell(value=str(raw or ""), confidence=0.4 if raw else 0.0, source_slug=slug)

    async def extract_page(
        self,
        *,
        slug: str,
        columns: list[ReportColumnDef],
    ) -> ExtractedRow | None:
        try:
            page = self.store.read_page(slug)
        except Exception:
            return None

        columns_spec = "\n".join(
            f"- key={col.key}: {col.label}. Instruction: {col.instruction}"
            for col in columns
        )

        cells: dict[str, ExtractedCell] = {}

        if not self.llm or not self.llm.available:
            raise ReportGenerationError(
                "Sorry, the AI report extractor is not connected. Please connect an LLM provider and try again."
            )
        try:
            raw = await self.llm.generate_json(
                safe_format(
                    EXTRACTION_PROMPT,
                    columns_spec=columns_spec,
                    title=page.meta.title,
                    content=trim_to_chars(page.content, 8000),
                ),
                temperature=0.05,
            )
        except Exception as exc:
            raise ReportGenerationError(
                "Sorry, the AI report extractor could not process this wiki page right now. "
                "Please try again or switch/check your LLM provider."
            ) from exc

        for col in columns:
            cells[col.key] = self._normalise_cell(raw.get(col.key, {}), slug=slug)

        return ExtractedRow(page_slug=slug, page_title=page.meta.title, cells=cells)

    async def extract_without_context(
        self,
        *,
        columns: list[ReportColumnDef],
    ) -> ExtractedRow:
        columns_spec = "\n".join(
            f"- key={col.key}: {col.label}. Instruction: {col.instruction}"
            for col in columns
        )

        cells: dict[str, ExtractedCell] = {}

        if not self.llm or not self.llm.available:
            raise ReportGenerationError(
                "Sorry, the AI report extractor is not connected. Please connect an LLM provider and try again."
            )
        try:
            raw = await self.llm.generate_json(
                safe_format(
                    KNOWLEDGE_EXTRACTION_PROMPT,
                    columns_spec=columns_spec,
                ),
                temperature=0.3,
            )
        except Exception as exc:
            raise ReportGenerationError(
                "Sorry, the AI report extractor could not process this report right now. "
                "Please try again or switch/check your LLM provider."
            ) from exc

        for col in columns:
            cells[col.key] = self._normalise_cell(raw.get(col.key, {}), slug="llm_knowledge")

        return ExtractedRow(page_slug="llm_knowledge", page_title="LLM Internal Knowledge", cells=cells)


# ---------------------------------------------------------------------------
# Exporters
# ---------------------------------------------------------------------------


def export_xlsx(rows: list[ExtractedRow], columns: list[ReportColumnDef]) -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Report"
    evidence_ws = wb.create_sheet("Evidence")

    # Header row
    headers = ["Page", "Title"] + [col.label for col in columns]
    header_fill = PatternFill("solid", fgColor="1E293B")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, vertical="center")

    # Data rows
    for row_idx, row in enumerate(rows, start=2):
        ws.cell(row=row_idx, column=1, value=row.page_slug)
        ws.cell(row=row_idx, column=2, value=row.page_title)
        for col_idx, col in enumerate(columns, start=3):
            cell_data = row.cells.get(col.key)
            val = cell_data.value if cell_data else ""
            ws.cell(row=row_idx, column=col_idx, value=val)

        # Zebra stripe
        if row_idx % 2 == 0:
            fill = PatternFill("solid", fgColor="F8FAFC")
            for col_idx in range(1, len(headers) + 1):
                ws.cell(row=row_idx, column=col_idx).fill = fill

    # Auto-width
    for col_cells in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 50)

    evidence_headers = ["Page", "Title", "Column", "Value", "Confidence", "Source Quote"]
    for col_idx, header in enumerate(evidence_headers, start=1):
        cell = evidence_ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    evidence_row = 2
    for row in rows:
        for col in columns:
            cell_data = row.cells.get(col.key)
            evidence_ws.cell(row=evidence_row, column=1, value=row.page_slug)
            evidence_ws.cell(row=evidence_row, column=2, value=row.page_title)
            evidence_ws.cell(row=evidence_row, column=3, value=col.label)
            evidence_ws.cell(row=evidence_row, column=4, value=cell_data.value if cell_data else "")
            evidence_ws.cell(row=evidence_row, column=5, value=cell_data.confidence if cell_data else 0.0)
            evidence_ws.cell(row=evidence_row, column=6, value=cell_data.quote if cell_data else "")
            evidence_ws.cell(row=evidence_row, column=6).alignment = Alignment(wrap_text=True)
            evidence_row += 1
    for col_cells in evidence_ws.columns:
        max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
        evidence_ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 70)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def is_value_empty_or_negative(cell: ExtractedCell | None) -> bool:
    if not cell:
        return True
    val = cell.value.strip()
    if not val:
        return True
    if cell.confidence == 0.0:
        return True
    
    # Check for common negative / not found sentences generated by LLMs
    val_lower = val.lower()
    negative_patterns = [
        "not found",
        "not mentioned",
        "no information",
        "not specified",
        "no educational",
        "no certification",
        "no qualification",
        "none mentioned",
        "none found",
        "does not mention",
        "is not mentioned",
        "are not mentioned",
        "not provided",
        "no mention",
    ]
    for pattern in negative_patterns:
        if pattern in val_lower and len(val) < 150:
            return True
    return False


def export_docx(
    rows: list[ExtractedRow],
    columns: list[ReportColumnDef],
    template_name: str = "Report",
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(template_name, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    for row in rows:
        # Check first if there is any content to render
        has_content = False
        for col in columns:
            cell_data = row.cells.get(col.key)
            if cell_data and not is_value_empty_or_negative(cell_data):
                has_content = True
                break

        if not has_content:
            continue

        # Heading for page
        h = doc.add_heading(row.page_title, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Render content
        for col in columns:
            cell_data = row.cells.get(col.key)
            if cell_data and not is_value_empty_or_negative(cell_data):
                val = cell_data.value
                if len(columns) > 1:
                    p_label = doc.add_paragraph()
                    run_label = p_label.add_run(f"{col.label}:")
                    run_label.bold = True
                
                add_markdown_to_docx(doc, val)
        
        doc.add_paragraph("") # separator paragraph

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(
    rows: list[ExtractedRow],
    columns: list[ReportColumnDef],
    template_name: str = "Report",
) -> bytes:
    import weasyprint

    sections_html = ""
    for row in rows:
        content_html = ""
        has_content = False
        for col in columns:
            cell_data = row.cells.get(col.key)
            if cell_data and not is_value_empty_or_negative(cell_data):
                val = cell_data.value
                if len(columns) > 1:
                    content_html += f"<p><strong>{escape(col.label)}:</strong></p>"
                content_html += markdown_to_html(val)
                has_content = True

        if has_content:
            sections_html += f"""
            <div class="page-section">
                <h2>{escape(row.page_title)}</h2>
                {content_html}
            </div>
            """

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 12px; color: #1e293b; margin: 20px; line-height: 1.6; }}
  h1 {{ color: #cc5a37; font-size: 24px; margin-bottom: 4px; border-bottom: 2px solid #cc5a37; padding-bottom: 6px; }}
  p.meta {{ color: #64748b; font-size: 11px; margin-bottom: 24px; }}
  h2 {{ color: #1e293b; font-size: 16px; margin-top: 20px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; }}
  .page-section {{ margin-bottom: 30px; }}
  strong {{ color: #cc5a37; }}
  table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 11px; }}
  th {{ background: #f1f5f9; color: #1e293b; border: 1px solid #cbd5e1; padding: 6px 10px; text-align: left; font-weight: 600; }}
  td {{ border: 1px solid #e2e8f0; padding: 6px 10px; }}
  ul, ol {{ margin: 8px 0; padding-left: 20px; }}
  li {{ margin-bottom: 4px; }}
  code {{ background: #f1f5f9; padding: 2px 4px; border-radius: 4px; font-family: monospace; font-size: 90%; }}
</style>
</head>
<body>
<h1>{escape(template_name)}</h1>
<p class="meta">Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}</p>
{sections_html}
</body>
</html>"""

    return weasyprint.HTML(string=html).write_pdf()


def markdown_to_html(text: str) -> str:
    import re
    # Normalize newlines
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = text.split('\n')
    
    html_out = []
    current_block_type = None # None, 'ul', 'ol', 'p', 'table'
    current_block_lines = []
    
    def close_block():
        nonlocal current_block_type, current_block_lines
        if not current_block_type or not current_block_lines:
            return
        
        if current_block_type == 'ul':
            items_html = "".join(f"<li>{parse_inline_markdown(line)}</li>" for line in current_block_lines)
            html_out.append(f"<ul>{items_html}</ul>")
        elif current_block_type == 'ol':
            items_html = "".join(f"<li>{parse_inline_markdown(line)}</li>" for line in current_block_lines)
            html_out.append(f"<ol>{items_html}</ol>")
        elif current_block_type == 'table':
            if len(current_block_lines) >= 2:
                hdr_cols = [c.strip() for c in current_block_lines[0].strip('|').split('|')]
                table_html = "<table><thead><tr>"
                for col in hdr_cols:
                    table_html += f"<th>{parse_inline_markdown(col)}</th>"
                table_html += "</tr></thead><tbody>"
                
                for row_line in current_block_lines[2:]:
                    row_cols = [c.strip() for c in row_line.strip('|').split('|')]
                    table_html += "<tr>"
                    for col in row_cols:
                        table_html += f"<td>{parse_inline_markdown(col)}</td>"
                    table_html += "</tr>"
                table_html += "</tbody></table>"
                html_out.append(table_html)
            else:
                fallback_text = "<br>".join(current_block_lines)
                html_out.append(f"<p>{parse_inline_markdown(fallback_text)}</p>")
        elif current_block_type == 'p':
            para_text = "<br>".join(current_block_lines)
            html_out.append(f"<p>{parse_inline_markdown(para_text)}</p>")
            
        current_block_type = None
        current_block_lines = []

    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            close_block()
            continue
            
        # Check if line is a header
        if line_strip.startswith('#'):
            close_block()
            level = len(re.match(r'^#+', line_strip).group(0))
            content = line_strip.lstrip('#').strip()
            h_level = min(max(level, 3), 5)
            html_out.append(f"<h{h_level}>{parse_inline_markdown(content)}</h{h_level}>")
            continue
            
        # Check if line is a bullet list item
        bullet_match = re.match(r'^([\-\*\+•])\s+(.*)', line_strip)
        if bullet_match:
            if current_block_type != 'ul':
                close_block()
                current_block_type = 'ul'
            current_block_lines.append(bullet_match.group(2))
            continue
            
        # Check if line is a numbered list item
        num_match = re.match(r'^(\d+)\.\s+(.*)', line_strip)
        if num_match:
            if current_block_type != 'ol':
                close_block()
                current_block_type = 'ol'
            current_block_lines.append(num_match.group(2))
            continue
            
        # Check if line is part of a table
        if '|' in line_strip:
            if current_block_type != 'table':
                close_block()
                current_block_type = 'table'
            current_block_lines.append(line_strip)
            continue
            
        # Otherwise, it's a normal paragraph line
        if current_block_type != 'p':
            close_block()
            current_block_type = 'p'
        current_block_lines.append(line_strip)
        
    close_block()
    return "\n".join(html_out)


def parse_inline_markdown(text: str) -> str:
    import re
    # Escape HTML tags first
    t = escape(text)
    # Bold: **bold**
    t = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', t)
    # Italic: *italic*
    t = re.sub(r'\*(.*?)\*', r'<em>\1</em>', t)
    # Code: `code`
    t = re.sub(r'`(.*?)`', r'<code>\1</code>', t)
    return t


def add_markdown_to_docx(doc, text: str) -> None:
    import re
    # Normalize newlines
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = text.split('\n')
    
    current_block_type = None
    current_block_lines = []
    
    def close_block():
        nonlocal current_block_type, current_block_lines
        if not current_block_type or not current_block_lines:
            return
            
        if current_block_type == 'ul':
            for line in current_block_lines:
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs_to_paragraph(p, line)
        elif current_block_type == 'ol':
            for line in current_block_lines:
                p = doc.add_paragraph(style='List Number')
                add_inline_runs_to_paragraph(p, line)
        elif current_block_type == 'table':
            if len(current_block_lines) >= 2:
                hdr_line = current_block_lines[0].strip().strip('|')
                hdr_cols = [c.strip() for c in hdr_line.split('|')]
                
                body_rows = []
                for line in current_block_lines[2:]:
                    row_line = line.strip().strip('|')
                    row_cols = [c.strip() for c in row_line.split('|')]
                    body_rows.append(row_cols)
                    
                num_cols = len(hdr_cols)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for col_idx, col_text in enumerate(hdr_cols):
                    p = hdr_cells[col_idx].paragraphs[0]
                    add_inline_runs_to_paragraph(p, col_text)
                    if p.runs:
                        p.runs[0].bold = True
                        
                # Body
                for row_data in body_rows:
                    row_cells = table.add_row().cells
                    for col_idx in range(min(num_cols, len(row_data))):
                        p = row_cells[col_idx].paragraphs[0]
                        add_inline_runs_to_paragraph(p, row_data[col_idx])
                doc.add_paragraph("")
            else:
                for line in current_block_lines:
                    p = doc.add_paragraph()
                    add_inline_runs_to_paragraph(p, line)
        elif current_block_type == 'p':
            p = doc.add_paragraph()
            add_inline_runs_to_paragraph(p, "\n".join(current_block_lines))
            
        current_block_type = None
        current_block_lines = []

    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            close_block()
            continue
            
        if line_strip.startswith('#'):
            close_block()
            level = len(re.match(r'^#+', line_strip).group(0))
            content = line_strip.lstrip('#').strip()
            doc.add_heading(content, level=min(level + 2, 4))
            continue
            
        bullet_match = re.match(r'^([\-\*\+•])\s+(.*)', line_strip)
        if bullet_match:
            if current_block_type != 'ul':
                close_block()
                current_block_type = 'ul'
            current_block_lines.append(bullet_match.group(2))
            continue
            
        num_match = re.match(r'^(\d+)\.\s+(.*)', line_strip)
        if num_match:
            if current_block_type != 'ol':
                close_block()
                current_block_type = 'ol'
            current_block_lines.append(num_match.group(2))
            continue
            
        if '|' in line_strip:
            if current_block_type != 'table':
                close_block()
                current_block_type = 'table'
            current_block_lines.append(line_strip)
            continue
            
        if current_block_type != 'p':
            close_block()
            current_block_type = 'p'
        current_block_lines.append(line_strip)
        
    close_block()


def add_inline_runs_to_paragraph(p, text: str) -> None:
    import re
    # Simple inline parser for **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            subparts = re.split(r'(\*.*?\*)', bold_text)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = p.add_run(subpart[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = p.add_run(subpart)
                    run.bold = True
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = p.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    p.add_run(subpart)


# ---------------------------------------------------------------------------
# Job Runner
# ---------------------------------------------------------------------------


class ReportRunner:
    def __init__(self, db: Session, store: WikiStore, llm: JsonLlm | None = None):
        self.db = db
        self.store = store
        self.llm = llm

    async def run(self, job: ReportJob) -> ReportJob:
        job.status = "processing"
        self.db.commit()

        try:
            template = self.db.get(ReportTemplate, job.template_id)
            if not template:
                raise ValueError(f"Template {job.template_id} not found")

            columns = [ReportColumnDef(**c) for c in json.loads(template.columns_json)]
            scope_slugs = json.loads(
                job.scope_slugs_json
                if job.scope_slugs_json is not None
                else (template.scope_slugs_json or "[]")
            )

            extractor = ReportExtractor(self.store, self.llm)
            rows: list[ExtractedRow] = []
            if scope_slugs:
                for slug in scope_slugs:
                    row = await extractor.extract_page(slug=slug, columns=columns)
                    if row:
                        rows.append(row)
            else:
                row = await extractor.extract_without_context(columns=columns)
                rows.append(row)

            if not rows:
                raise ReportGenerationError(
                    "No selected wiki pages could be read for this report. Please update the template scope and try again."
                )

            job.results_json = json.dumps([r.model_dump() for r in rows])

            # Export
            fmt = job.export_format
            export_name = f"report_{job.id[:8]}.{fmt}"
            export_dir = self.store.root / "reports"
            export_dir.mkdir(parents=True, exist_ok=True)
            export_path = export_dir / export_name

            if fmt == "xlsx":
                export_path.write_bytes(export_xlsx(rows, columns))
            elif fmt == "docx":
                export_path.write_bytes(export_docx(rows, columns, template.name))
            elif fmt == "pdf":
                export_path.write_bytes(export_pdf(rows, columns, template.name))

            job.file_path = str(export_path)
            job.status = "done"
            job.completed_at = datetime.now(UTC)

        except Exception as exc:
            job.status = "failed"
            job.error_message = str(exc)
            job.completed_at = datetime.now(UTC)

        self.db.commit()
        return job

    def create_job(
        self,
        *,
        workspace_id: str,
        template_id: str,
        export_format: str,
        created_by: str | None = None,
    ) -> ReportJob:
        job = ReportJob(
            id=str(uuid.uuid4()),
            workspace_id=workspace_id,
            template_id=template_id,
            export_format=export_format,
            created_by=created_by,
        )
        self.db.add(job)
        self.db.commit()
        self.db.refresh(job)
        return job


class ReportGenerationError(RuntimeError):
    pass


async def generate_report_from_chat(
    db: Session,
    user: Any,
    workspace: Any,
    request: Any,
) -> Any:
    from app.services.llm_factory import build_user_llm
    from app.api.deps import wiki_store_for_workspace
    from app.llmwiki.prompts import ANALYZE_CHAT_REPORT_PROMPT
    from app.schemas.llmwiki import ChatResponse, AgentTrace
    import json
    import uuid

    llm = build_user_llm(db, user)
    if not llm or not llm.available:
        from app.llmwiki.groq import GroqClient
        llm = GroqClient()

    if not llm or not llm.available:
        return ChatResponse(
            session_id=request.session_id,
            answer="I am sorry, but neither your personal LLM provider nor the system LLM provider is configured. Please configure an LLM provider to continue.",
            route="direct",
            difficulty="easy",
        )

    store = wiki_store_for_workspace(workspace)
    pages = store.list_pages()
    if not pages:
        return ChatResponse(
            session_id=request.session_id,
            answer="I am sorry, there are no wiki pages in this workspace to generate a report from. Please upload a PDF or create a wiki page first.",
            route="direct",
            difficulty="easy",
        )

    wiki_pages_list = "\n".join(
        f"- slug: {p.slug}, title: {p.title}, summary: {p.summary}"
        for p in pages
    )
    prompt = safe_format(
        ANALYZE_CHAT_REPORT_PROMPT,
        wiki_pages_list=wiki_pages_list,
        question=request.question,
    )

    try:
        analysis = await llm.generate_json(prompt, temperature=0.1)
    except Exception as exc:
        print(f"[Error] generate_report_from_chat analysis failed: {exc}")
        return ChatResponse(
            session_id=request.session_id,
            answer="I'm sorry, but I am currently unable to process your report request. Please try again shortly.",
            route="direct",
            difficulty="easy",
        )

    # Validate output
    name = (analysis.get("name") or "").strip()
    if not name:
        name = f"Chat Report - {datetime.now(UTC).strftime('%Y-%m-%d %H:%M')}"
    
    desc = (analysis.get("description") or "").strip()
    if not desc:
        desc = f"Generated dynamically from chat request: '{request.question}'"

    fmt = (analysis.get("export_format") or "xlsx").strip().lower()
    if fmt not in {"pdf", "xlsx", "docx"}:
        fmt = "xlsx"

    cols = analysis.get("columns") or []
    if not cols:
        cols = [{"key": "summary", "label": "Summary", "instruction": "Provide a brief summary of this page."}]
    
    sections = analysis.get("sections") or []

    scope_slugs = analysis.get("scope_slugs") or []
    if not scope_slugs:
        scope_slugs = [p.slug for p in pages]

    # Create the template
    template = ReportTemplate(
        id=str(uuid.uuid4()),
        workspace_id=workspace.id,
        name=name,
        description=desc,
        columns_json=json.dumps(cols),
        sections_json=json.dumps(sections),
        scope_slugs_json=json.dumps(scope_slugs),
        created_by=user.id,
    )
    db.add(template)
    db.commit()
    db.refresh(template)

    # Create the job
    job = ReportJob(
        id=str(uuid.uuid4()),
        workspace_id=workspace.id,
        template_id=template.id,
        export_format=fmt,
        created_by=user.id,
        scope_slugs_json=template.scope_slugs_json,
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    # Run the report job
    runner = ReportRunner(db, store, llm)
    try:
        await runner.run(job)
    except Exception as exc:
        job.status = "failed"
        job.error_message = str(exc)
        db.commit()

    if job.status == "done":
        used_slugs_str = ", ".join(f"`{slug}`" for slug in scope_slugs)
        ans = (
            f"I have successfully generated the report **{template.name}**.\n\n"
            f"📁 **File Details**:\n"
            f"- **Format**: {fmt.upper()}\n"
            f"- **Scope**: {used_slugs_str}\n\n"
            f"📥 **Download Link**:\n"
            f"[Click here to download your report](/api/v1/reports/{job.id}/download?format={fmt})\n\n"
            f"You can also manage this template and view all past generation jobs in the **📊 Report Generator** modal by clicking the report icon at the bottom of the sidebar."
        )
    else:
        ans = "I'm sorry, but I am currently unable to process your report request. Please try again shortly."

    return ChatResponse(
        session_id=request.session_id,
        answer=ans,
        route="wiki",
        difficulty="hard",
        citations=[],
        used_pages=scope_slugs,
        knowledge_gap_created=(job.status != "done"),
        agent_trace=[AgentTrace(
            agent="chat_report_generator",
            action="generate_report",
            confidence=1.0,
            notes=f"Created template '{template.name}' and job {job.id}."
        )]
    )
